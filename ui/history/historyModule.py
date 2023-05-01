import tkinter as Tk
from tkinter import ttk
import matplotlib
import matplotlib.pyplot as plt
matplotlib.use('TKAgg')
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.figure import Figure
import os
from i18n import _
from keyboard.keyboard import KeyBoard
from image.image import ImageAdrr
import numpy as np
from typing import TYPE_CHECKING
if TYPE_CHECKING:
    from main import Application
from tkinter import scrolledtext
current_directory = os.path.dirname(os.path.realpath(__file__))
parent_directory = os.path.dirname(os.path.dirname(current_directory))
import ctypes
from numpy.ctypeslib import ndpointer
ad7609 = ctypes.CDLL(f'{parent_directory}/ad7609BTZ.so')
from digitalFilter.digitalFilter import filter_data
from Calculation.calculate import *
import PlotData.PlotData as Pd
import pms.popMessage as pms
import sqlite3 as lite
import fileOperation.fileOperation as file_operation
import report.report as rp
import qrcode
balancingOrder=0
click_stop_flag=False
blink=0
blink1=0
track_flag=0
confirm_flag=0
def testVal(inStr, acttyp):
    if acttyp == '1':  # insert
        if not inStr.isdigit():
            return False
    return True

class History(Tk.Frame):
    def __init__(self, parent: "Application"):
        self.parent = parent
        imageAddress = ImageAdrr()
        self.homePhoto = imageAddress.homePhoto
        self.arrowPhoto = imageAddress.arrowPhoto
        self.style = ttk.Style()
        self.style.configure('normal.TButton', font=('Chakra Petch', 15), borderwidth=5, justify=Tk.CENTER)
        self.style.configure('custom.Accent.TButton', font=('Chakra Petch', 10), bordercolor='black', borderwidth=4,
                               justify=Tk.CENTER)
        self.style.configure('feature.Accent.TButton', font=('Chakra Petch', 15), borderwidth=1, justify=Tk.CENTER)
        self.style.configure('normal.TLabel', font=('Chakra Petch', 13), background='white')
        self.style.configure('red.TLabel', font=('Chakra Petch', 13), background='white', foreground='#C40069')
        self.con = lite.connect(f'{parent_directory}/company.db')
        self.mainFrame = Tk.Frame(self.parent, bd=1, bg='white', width=1024, height=600)
        self.mainFrame.pack(side=Tk.TOP, fill=Tk.BOTH, expand=1)
        self.mainFrame.pack_propagate(0)

        self.featureFrame = Tk.Frame(self.mainFrame, bd=1, bg='white', width=1024, height=80)
        self.featureFrame.pack(side=Tk.TOP, fill=Tk.BOTH, expand=1)
        self.featureFrame.pack_propagate(0)

        self.creat_setting_feature_panel()
        self.historyConfigFrame = historyConfig(self.mainFrame, self.infoLabel2, self.con, self.parent.origin_config.history_config_struct)
        self.historyConfigFrame.pack(side=Tk.TOP, fill=Tk.X, expand=1)
        self.historyConfigFrame.pack_propagate(0)

        self.historyAnalysisFrame = historyAnalysis(self.mainFrame, self.infoLabel2, self.con, self.parent.origin_config.history_config_struct)
        self.historyAnalysisFrame.pack(side=Tk.TOP, fill=Tk.BOTH, expand=1)
        self.historyAnalysisFrame.pack_propagate(0)
        self.historyAnalysisFrame.pack_forget()

        self.bearingFrequencyFrame = bearingFrequency(self.mainFrame, self.infoLabel2)
        self.bearingFrequencyFrame.pack(side=Tk.TOP, fill=Tk.BOTH, expand=1)
        self.bearingFrequencyFrame.pack_propagate(0)
        self.bearingFrequencyFrame.pack_forget()

        self.parent.bind_class('TEntry', "<FocusIn>", self.show_key_board)
        self.parent.bind_class('TCombobox', "<<ComboboxSelected>>", self.change_state)

    def show_key_board(self, event):
        global confirm_flag
        confirm_flag=0
        self.historyConfigFrame.applyBt.configure(state="normal")
        self.widget = self.get_focus_widget()
        self.keyboardFrame = KeyBoard(self.widget)
        parentName = event.widget.winfo_parent()
        self.parent1 = event.widget._nametowidget(parentName)
        self.parent1.focus()

    def get_focus_widget(self):
        widget = self.parent.focus_get()
        return widget

    def change_state(self, event):
        global confirm_flag
        confirm_flag=0
        self.historyConfigFrame.applyBt.configure(state="normal")

    def creat_setting_feature_panel(self):

        self.homeBt = ttk.Button(self.featureFrame, style='normal.TButton', text=_("Home"), image=self.homePhoto,
                                 compound=Tk.TOP,
                                 command=self.go_home)
        self.homeBt.place(relx=0.0, rely=0.018, width=100, height=72)
        self.homeBt.image = self.homePhoto

        barrie=Tk.Frame(self.featureFrame, width=3, height=72, background='grey')
        barrie.place(relx=0.11, rely=0.018)

        self.configBt = ttk.Button(self.featureFrame, style='feature.Accent.TButton', text=_("Config"),
                                    command=self.on_config_button_clicked)
        self.configBt.place(relx=0.122, rely=0.018, width=115, height=72)

        self.arrowLabel = ttk.Label(self.featureFrame, style='normal.TLabel', image=self.arrowPhoto)
        self.arrowLabel.place(relx=0.237, rely=0.25)
        self.arrowLabel.image = self.arrowPhoto
        
        self.analysisBt = ttk.Button(self.featureFrame, style='normal.TButton', text=_("History\nAnalysis"),
                                   command=self.on_analysis_button_clicked)
        self.analysisBt.place(relx=0.265, rely=0.018, width=115, height=72)

        self.bearingBt = ttk.Button(self.featureFrame, style='normal.TButton', text=_("Bearing\nFrequency"),
                                   command=self.on_bearing_button_clicked)
        self.bearingBt.place(relx=0.385, rely=0.018, width=125, height=72)


        self.infoFrame= Tk.Frame(self.featureFrame, width=451, height=72, bg='white', bd=0)
        self.infoFrame.place(relx=0.52, rely=0.018)

        self.infoLabel1=ttk.Label(self.infoFrame, text=_("Information"), style="red.TLabel")
        self.infoLabel1.grid(column=0, row=0, padx=0, pady=5, sticky='w')

        self.infoLabel2 = ttk.Label(self.infoFrame, text="OK.", style="normal.TLabel", width=42)
        self.infoLabel2.grid(column=0, row=1, padx=0, pady=5, sticky='w')

    def on_analysis_button_clicked(self):
        global confirm_flag
        if confirm_flag==1:
            self.historyConfigFrame.pack_forget()
            self.bearingFrequencyFrame.pack_forget()
            self.historyAnalysisFrame.pack(side=Tk.TOP, fill=Tk.BOTH, expand=1)
            self.configBt.configure(style="normal.TButton")
            self.analysisBt.configure(style="feature.Accent.TButton")
            self.bearingBt.configure(style="normal.TButton")
        else:
            self.infoLabel2.configure(text=_("Click APPLY button before using this funtion."))

    def on_config_button_clicked(self):
        self.historyConfigFrame.pack(side=Tk.TOP, fill=Tk.BOTH, expand=1)
        self.historyAnalysisFrame.pack_forget()
        self.bearingFrequencyFrame.pack_forget()
        self.configBt.configure(style="feature.Accent.TButton")
        self.analysisBt.configure(style="normal.TButton")
        self.bearingBt.configure(style="normal.TButton")

    def on_bearing_button_clicked(self):
        self.bearingFrequencyFrame.pack(side=Tk.TOP, fill=Tk.BOTH, expand=1)
        self.historyAnalysisFrame.pack_forget()
        self.historyConfigFrame.pack_forget()
        self.configBt.configure(style="normal.TButton")
        self.analysisBt.configure(style="normal.TButton")
        self.bearingBt.configure(style="feature.Accent.TButton")

    def go_home(self):
        self.mainFrame.destroy()
        self.parent.go_to_home_page()

class historyConfig(Tk.Frame):
    def __init__(self, parent, infoLabel, db_connect, history_config_struct):
        super().__init__(parent, width=1024, height=520, background='white')
        self.parent=parent
        self.infoLabel=infoLabel
        self.prjDate=Tk.StringVar()
        self.con=db_connect
        self.style = ttk.Style()
        self.style.configure('history.TLabel', font=('Chakra Petch', 13), bg='white')
        self.style.configure('history.TLabelframe', font=('Chakra Petch', 15), bg='white', borderwidth=0)
        self.style.configure('history.TButton', font=('Chakra Petch', 15))
        self.style.configure('history.TEntry', font=('Chakra Petch', 15))
        self.style.configure('history.Switch.TCheckbutton', font=('Chakra Petch', 13))
        self.style.configure("Custom.Treeview", rowheight=30, font=('Chakra Petch', 12))
        self.creatConfig(history_config_struct)

    def creatConfig(self, history_config_struct):
        self.historyParam1=Tk.StringVar()
        self.historyParam2=Tk.StringVar()
        self.historyParam3=Tk.StringVar()
        self.historyParam4=Tk.StringVar()
        self.historyParam5=Tk.StringVar()
        self.historyParam6=Tk.StringVar()
        self.historyParam7=Tk.StringVar()
        self.historyParam8=Tk.StringVar()
        self.tsa_var = Tk.IntVar()
        self.var_1 = Tk.BooleanVar(self, True)
        self.var_2 = Tk.BooleanVar(self, True)
        self.var_3 = Tk.BooleanVar(self, True)
        self.var_4 = Tk.BooleanVar(self, True)

        self.historyParam1.set(history_config_struct["ProjectID"])
        self.historyParam2.set(history_config_struct["SensorPosition"])
        self.historyParam3.set(history_config_struct["PlotType"])
        self.historyParam4.set(history_config_struct["FilterFrom"])
        self.historyParam5.set(history_config_struct["FilterTo"])
        self.historyParam6.set(history_config_struct["ViewLimit"])
        self.historyParam7.set(history_config_struct["TrackingResolution"])
        self.historyParam8.set(history_config_struct["TsaBin"])
        self.tsa_var.set(history_config_struct["TSA"])


        historyFrame = ttk.LabelFrame(self, text=_('Configuration'), style='history.TLabelframe')
        historyFrame.pack(side="left", fill='y')

        projectIDLabel = ttk.Label(historyFrame, text=_('Project Code'), style="history.TLabel")
        projectIDLabel.grid(column=0, row=0, padx=5, pady=4, sticky="w")
        projectIDEntry=ttk.Entry(historyFrame, width=14, textvariable=self.historyParam1, font=('Chakra Petch', 13),
                                validate="key")
        projectIDEntry.grid(column=1, row=0, padx=(0,10), pady=4, sticky="e")

        positionLabel = ttk.Label(historyFrame, text=_('Sensor Position'), style="history.TLabel")
        positionLabel.grid(column=0, row=1, padx=5, pady=4, sticky="w")
        self.positionCombo=ttk.Combobox(historyFrame, width=10, textvariable=self.historyParam2, state="readonly",
                                    font=('Chakra Petch', 13))
        self.positionCombo.bind("<Button-1>", self.find_position_list)
        self.positionCombo.grid(column=1, row=1, padx=(0, 10), pady=4, ipadx=7, sticky="e")

        filterFromLabel = ttk.Label(historyFrame, text=_("Bandpass Filter From"), style="history.TLabel")
        filterFromLabel.grid(column=0, row=3, padx=5, pady=4, sticky='w')
        filterFromEntry=ttk.Entry(historyFrame, width=14, textvariable=self.historyParam4, font=('Chakra Petch', 13),
                                validate="key")
        filterFromEntry.grid(column=1, row=3, padx=(0,10), pady=4, sticky='e')

        filterToLabel = ttk.Label(historyFrame, text=_("Bandpass Filter To"), style="history.TLabel")
        filterToLabel.grid(column=0, row=4, padx=5, pady=4, sticky='w')
        filterToEntry=ttk.Entry(historyFrame, width=14, textvariable=self.historyParam5, font=('Chakra Petch', 13),
                                validate="key")
        filterToEntry.grid(column=1, row=4, padx=(0, 10), pady=4, sticky='e')

        viewLabel = ttk.Label(historyFrame, text=_("View Limit"), style="history.TLabel")
        viewLabel.grid(column=0, row=5, padx=5, pady=4, sticky='w')
        viewEntry=ttk.Entry(historyFrame, width=14, textvariable=self.historyParam6, font=('Chakra Petch', 13),
                                validate="key")
        viewEntry.grid(column=1, row=5, padx=(0, 10), pady=4, sticky='e')

        meshLabel = ttk.Label(historyFrame, text=_("Tracking speed"), style="history.TLabel")
        meshLabel.grid(column=0, row=6, padx=5, pady=4, sticky='w')
        meshEntry=ttk.Entry(historyFrame, width=14, textvariable=self.historyParam7, font=('Chakra Petch', 13),
                                validate="key")
        meshEntry.grid(column=1, row=6, padx=(0, 10), pady=4, sticky='e')
        
        tsaCheckButton=ttk.Checkbutton(historyFrame, text=_("Use Average"), offvalue=0, onvalue=1, 
                            variable=self.tsa_var, command=self.update_text_tsa, style="history.Switch.TCheckbutton")
        tsaCheckButton.grid(column=1,row=7, padx=(0,5), pady=2, sticky='e')

        tsaLabel = ttk.Label(historyFrame, text=_("Average Bin"), style="history.TLabel")
        tsaLabel.grid(column=0, row=8, padx=5, pady=4, sticky='w')
        self.tsaBin=ttk.Combobox(historyFrame, width=10, textvariable=self.historyParam8, state="readonly",
                                    font=('Chakra Petch', 13))
        self.tsaBin['value'] = ('2048','4096','8192','16384','32768')
        self.tsaBin.grid(column=1, row=8, padx=(0, 10), pady=4, ipadx=7, sticky='e')

        checkbox_1 = ttk.Checkbutton(historyFrame, text="Show Vel RMS", variable=self.var_1, command=self.checkbox_clicked)
        checkbox_1.grid(row=9, column=0, padx=5, pady=4, sticky="w")

        checkbox_2 = ttk.Checkbutton(historyFrame, text="Show Bearing gE", variable=self.var_2, command=self.checkbox_clicked)
        checkbox_2.grid(row=9, column=1, padx=5, pady=4, sticky="w")

        checkbox_3 = ttk.Checkbutton(historyFrame, text="Show Gear Acc", variable=self.var_3, command=self.checkbox_clicked)
        checkbox_3.grid(row=10, column=0, padx=5, pady=4, sticky="w")

        checkbox_4 = ttk.Checkbutton(historyFrame, text="Show Bearing HFCF", variable=self.var_4, command=self.checkbox_clicked)
        checkbox_4.grid(row=10, column=1, padx=5, pady=4, sticky="w")

        self.applyBt = ttk.Button(historyFrame, style='Accent.TButton', text=_("APPLY"),
                                   command=lambda: self.update_config_struct(history_config_struct))
        self.applyBt.grid(column=0, row=11, padx=(5, 5), pady=4, ipadx=50, ipady=4, sticky='w')

        databaseConfigFrame = ttk.LabelFrame(self, text=(''), style="history.TLabelframe")
        databaseConfigFrame.pack(side="left", fill='y')
        deletePrjButton = ttk.Button(historyFrame, text=_("Delete Project"), style='Accent.TButton',
                            command=self.delete_project)
        deletePrjButton.grid(column=1, row=11, padx=(5, 5), pady=4, ipadx=20, ipady=4, sticky='w')

        self.scrollbar = ttk.Scrollbar(self)
        self.scrollbar.pack(side="right", fill="y")

        self.PrjTable = ttk.Treeview(self, yscrollcommand=self.scrollbar.set, show=("tree",), style="Custom.Treeview")
        self.PrjTable.bind('<<TreeviewSelect>>', self.get_selected_cell)

        self.PrjTable.pack(side="left", expand=True, fill="both")
        self.PrjTable['columns'] = ('PrjID', 'Date', 'Pos', 'SampleRate')
        self.PrjTable.column("#0", anchor='w', width=120)
        self.PrjTable.column("PrjID",anchor='w', width=120)
        self.PrjTable.column("Date",anchor=Tk.CENTER,width=120)
        self.PrjTable.column("Pos",anchor=Tk.CENTER,width=60)
        self.PrjTable.column("SampleRate",anchor=Tk.CENTER,width=80)

        self.PrjTable.heading("#0",text="",anchor=Tk.CENTER)
        self.PrjTable.heading("PrjID",text="ID",anchor=Tk.CENTER)
        self.PrjTable.heading("Date",text=_("Date"),anchor=Tk.CENTER)
        self.PrjTable.heading("Pos",text=_("Position"),anchor=Tk.CENTER)
        self.PrjTable.heading("SampleRate",text=_("Sample rate"),anchor=Tk.CENTER)
        self.scrollbar.config(command=self.PrjTable.yview)
        prjCodeArr=self.load_all_project_code()
        for i in range(len(prjCodeArr)):
            data_arr=self.load_data_base_on_code(prjCodeArr[i])
            companyName=self.find_company_name_base_on_code(prjCodeArr[i])
            date_arr=[arr[0] for arr in data_arr]
            pos_arr=[arr[1] for arr in data_arr]
            sample_rate_arr=[arr[2] for arr in data_arr]
            self.PrjTable.insert(parent='', index='end', iid=i, text=prjCodeArr[i], values=(str(companyName), '', '', ''))
            for k in range(len(pos_arr)):
                self.PrjTable.insert(parent=i,index='end',text='', values=(str(prjCodeArr[i]),str(date_arr[k]), \
                            str(pos_arr[k]), str(sample_rate_arr[k])))
        # self.PrjTable.item(open=True)
        # print("load_prj_code:", load_data_arr)

        # for i in range(len(load_data_arr)):
        #                 self.PrjTable.insert(parent='',index='end',iid=i,text='', values=(str(load_data_arr[i][0]),str(load_data_arr[i][1]), \
        #                                     str(load_data_arr[i][2]), str(load_data_arr[i][3])))

        
    def checkbox_clicked(self):
        self.applyBt.configure(state='normal')
        
    def get_selected_cell(self, event):
        global confirm_flag
        confirm_flag=0
        self.applyBt.configure(state='normal')
        
        try:
            selected_item = self.PrjTable.focus() # get the selected item
            self.PrjTable.item(selected_item, open=True)
            if self.PrjTable.parent(selected_item) !='':
                prjID = self.PrjTable.item(selected_item)['values'][0] # get the value of the cell
                position=self.PrjTable.item(selected_item)['values'][2]
                date=self.PrjTable.item(selected_item)['values'][1]
                self.historyParam1.set(prjID)
                self.historyParam2.set(position)
                self.prjDate.set(date)
            else:
                prjID = self.PrjTable.item(selected_item)['text']
                self.historyParam1.set(prjID)
                self.historyParam2.set('')
                for item in self.PrjTable.get_children():
                    if item != selected_item:
                        self.PrjTable.item(item, open=False)
        except:
            pass

    def update_text_tsa(self):
        global confirm_flag
        confirm_flag=0
        self.applyBt.configure(state='normal')
        txt=self.tsa_var.get()
        if txt==1:
            self.tsaBin.configure(state='normal')
        else:
            self.tsaBin.configure(state='disabled')

    def update_config_struct(self, history_config_struct):
        global confirm_flag
        confirm_flag=1
        history_config_struct["ProjectID"]=self.historyParam1.get()
        history_config_struct["SensorPosition"]=self.historyParam2.get()
        history_config_struct["PlotType"]=self.historyParam3.get()
        history_config_struct["TsaBin"]=int(self.historyParam8.get())
        history_config_struct["TSA"]=self.tsa_var.get()
        history_config_struct["Rms"]=self.var_1.get()
        history_config_struct["A_Pk"]=self.var_3.get()
        history_config_struct["gE"]=self.var_2.get()
        history_config_struct["HFCF"]=self.var_4.get()
        tempViewValue=  self.historyParam6.get()
        tempMeshValue=  self.historyParam7.get()
        tempFilterFrom= self.historyParam4.get()
        tempFilterTo=   self.historyParam5.get()
        if is_number(tempMeshValue)==False or is_number(tempViewValue)==False or is_number(tempFilterFrom)==False or \
            is_number(tempFilterTo)==False:
            self.infoLabel.config(text=_("Parameter errors. Check the input type."))
        else:
            history_config_struct["TrackingResolution"]=int(tempMeshValue)
            history_config_struct["ViewLimit"]=int(tempViewValue)
            history_config_struct["FilterFrom"]=int(tempFilterFrom)
            history_config_struct["FilterTo"]=int(tempFilterTo)
            
            self.infoLabel.config(text=_("OK")) 
            self.applyBt.configure(state='disable')
    def load_project(self):
        with self.con:
            cur=self.con.cursor()
            cur.execute(f"SELECT CODE, DATE, POS, Sample_rate FROM DATA ORDER BY DATE ASC;")
            load_data = cur.fetchall()
            load_data_arr = [i for i in load_data]
            return load_data_arr

    def load_all_project_code(self):
        with self.con:
            cur=self.con.cursor()
            cur.execute(f"SELECT DISTINCT CODE FROM DATA ORDER BY DATE ASC;")
            load_data = cur.fetchall()
            load_data_arr = [i for i in load_data]
            return [arr[0] for arr in load_data_arr]

    def load_data_base_on_code(self, prjCode):
        with self.con:
            cur=self.con.cursor()
            cur.execute(f"SELECT DATE, POS, Sample_rate FROM DATA WHERE CODE=? ORDER BY DATE ASC", (prjCode,))
            load_data = cur.fetchall()
            load_data_arr = [i for i in load_data]
            return load_data_arr
    
    def find_company_name_base_on_code(self, prjCode):
        with self.con:
            cur=self.con.cursor()
            cur.execute(f"SELECT COM_ID FROM Project_ID WHERE CODE=?", (prjCode,))
            load_data=cur.fetchall()
            com_id=load_data[0][0]
            cur.execute(f"SELECT NAME FROM Company_ID WHERE COM_ID =?", (com_id,))
            name=cur.fetchall()
            return name[0][0]
        
    def find_position_list(self, event):
        prjCode=self.historyParam1.get()
        self.historyParam2.set("")
        with self.con:
            cur=self.con.cursor()
            cur.execute(f"SELECT POS FROM DATA WHERE CODE = ? ORDER BY DATE ASC", (prjCode,))
            load_data = cur.fetchall()
            load_data_arr = [i for i in load_data]
        posList=[]
        for i in range(len(load_data)):
            posList.append(load_data_arr[i][0])
        new_list = []
        [new_list.append(item) for item in posList if item not in new_list]
        self.positionCombo['value'] = tuple(new_list)
        
    def delete_project(self):
        try:
            prjCode = self.historyParam1.get()
            position =  self.historyParam2.get().replace(' ', '')
            date = self.prjDate.get()
            if prjCode!='':
                with self.con:
                    cur=self.con.cursor()
                    if pms.general_warning(_("Do you want to delete this project?"))==True:
                        if position=='':
                            cur.execute(f"""DELETE FROM DATA WHERE CODE = ? """, (prjCode,))
                            cur.execute(f"""DELETE FROM Project_ID WHERE CODE = ? """, (prjCode,))
                        else:
                            cur.execute(f"""DELETE FROM DATA WHERE CODE = ? AND DATE = ? AND POS = ?    """, (prjCode, date, position,))
                        self.infoLabel.configure(text=_("Selected project is deleted."))
                        for row in self.PrjTable.get_children():
                            self.PrjTable.delete(row)
                        prjCodeArr=self.load_all_project_code()
                        for i in range(len(prjCodeArr)):
                            data_arr=self.load_data_base_on_code(prjCodeArr[i])
                            companyName=self.find_company_name_base_on_code(prjCodeArr[i])
                            date_arr=[arr[0] for arr in data_arr]
                            pos_arr=[arr[1] for arr in data_arr]
                            sample_rate_arr=[arr[2] for arr in data_arr]
                            self.PrjTable.insert(parent='', index='end', iid=i, text=prjCodeArr[i], values=(str(companyName), '', '', ''))
                            for k in range(len(pos_arr)):
                                self.PrjTable.insert(parent=i,index='end',text='', values=(str(prjCodeArr[i]),str(date_arr[k]), \
                                            str(pos_arr[k]), str(sample_rate_arr[k])))
                    else:
                        pass
            else:
                pms.empty_entry_error(_("Project Code"))
        except:
            self.infoLabel.configure(text=_("Delete project errors"))

class historyAnalysis(Tk.Frame):
    def __init__(self, parent:"mainFrame", infoLabel, db_connect, history_config_struct):
        super().__init__(parent, width=1008 , height=504 , bg='white')
        self.parent = parent
        self.history_config_struct = history_config_struct
        self.infoLabel = infoLabel
        self.con=db_connect
        self.history_analysis_page()
    def history_analysis_page(self):
        self.historyPlotFrame = HistoryPlotCanvas(self)
        self.historyPlotFrame.pack(side=Tk.LEFT, fill=Tk.BOTH, expand=1)
        self.historyPlotFrame.pack_propagate(0)

        self.historySideFrame = SideButtonFrame(self, self.history_config_struct, self.con, self.infoLabel,
                                                 self.historyPlotFrame.canvas4)
        self.historySideFrame.pack(side=Tk.RIGHT, fill=Tk.Y, expand=1)
        self.historySideFrame.pack_propagate(0)

class bearingFrequency(Tk.Frame):
    def __init__(self, parent:"mainFrame", infoLabel):
        super().__init__(parent, width=1008 , height=504 , bg='white')
        self.parent = parent
        self.infoLabel = infoLabel
        self.con = lite.connect(f'{parent_directory}/brgfreq.db')
        
        self.bearingDetailConfig()
    
    def bearingDetailConfig(self):
        self.brgParam1 = Tk.StringVar()
        self.brgParam2 = Tk.StringVar()

        self.brgParam2.set(1500)
        parentFrame = Tk.LabelFrame(self, text='', font=('Chakra Petch', 13), border=0, bg='white')
        parentFrame.pack(side=Tk.TOP, fill=Tk.BOTH)
        imputFrame = ttk.LabelFrame(parentFrame, text=_('Input'))
        imputFrame.grid(column=0, row=0, padx=5, pady=0, sticky='wn')
        imageFrame = ttk.LabelFrame(parentFrame, text=(''))
        imageFrame.grid(column=0, row=1, padx=5, pady=0, sticky='wn')
        tableFrame = ttk.LabelFrame(parentFrame, text=_('Bearing frequency table'))
        tableFrame.grid(column=2, row=0, rowspan=2, padx=5, pady=0, sticky='wn')
        
        imageAdress=ImageAdrr()
        image = imageAdress.bearingPhoto
        self.fig5 = Figure(figsize=(4,3))
        self.ax_41 = self.fig5.add_subplot()
        self.ax_41.set_position([0.01, 0.01, 1, 1])
        self.ax_41.set_xticks([])
        self.ax_41.set_yticks([])
        self.ax_41.imshow(image)
        self.ax_41.axis('off')
        self.ax_41.set_visible(True)
        self.fig5.set_visible(True)
        self.canvas4 = FigureCanvasTkAgg(self.fig5, master=imageFrame)
        self.canvas4.get_tk_widget().pack(side=Tk.TOP, fill=Tk.BOTH, expand=1)

        bearingNameLabel = ttk.Label(imputFrame, text=_('Bearing Name'), style='config.TLabel')
        bearingNameLabel.grid(column=0, row=0, padx=(10, 10), pady=5, sticky='w')

        self.nameEntry = ttk.Entry(imputFrame, width=10, textvariable=self.brgParam1,
                                  state="normal", font=('Chakra Petch', 13))
        self.nameEntry.grid(column=1, row=0, padx=(10, 10), pady=5, ipadx=3, sticky='e')  

        speedLabel = ttk.Label(imputFrame, text=_("Speed(RPM)"), style='config.TLabel')
        speedLabel.grid(column=0, row=1, padx=(10, 10), pady=5, sticky='w')
        self.speedEntry = ttk.Entry(imputFrame, width=10, textvariable=self.brgParam2,
                                  state="normal", font=('Chakra Petch', 13))
        self.speedEntry.grid(column=1, row=1, padx=(10, 10), pady=5, ipadx=3, sticky='e')

        self.searchBt = ttk.Button(imputFrame, text=_("SEARCH"), style='Accent.TButton',
                                 command=self.on_search_button_clicked)
        self.searchBt.grid(column=2, row=3, padx=5, pady=5, ipadx=10, sticky='e')

        self.myTable = ttk.Treeview(tableFrame)
        self.myTable.grid(column=0, row=0, padx=5, pady=5, sticky='e')
        self.myTable['columns'] = ('name', 'BPFO', 'BPFI', 'BSF', 'FTF')
        self.myTable.column("#0", width=0,  stretch=Tk.NO)
        self.myTable.column("name",anchor=Tk.CENTER, width=100)
        self.myTable.column("BPFO",anchor=Tk.CENTER,width=100)
        self.myTable.column("BPFI",anchor=Tk.CENTER,width=100)
        self.myTable.column("BSF",anchor=Tk.CENTER,width=100)
        self.myTable.column("FTF",anchor=Tk.CENTER,width=100)
        self.myTable.heading("#0",text="",anchor=Tk.CENTER)
        self.myTable.heading("name",text="Name",anchor=Tk.CENTER)
        self.myTable.heading("BPFO",text="BPFO",anchor=Tk.CENTER)
        self.myTable.heading("BPFI",text="BPFI",anchor=Tk.CENTER)
        self.myTable.heading("BSF",text="BSF",anchor=Tk.CENTER)
        self.myTable.heading("FTF",text="FTF",anchor=Tk.CENTER)

    def on_search_button_clicked(self):
        global confirm_flag
        confirm_flag=1
        try:
            bearingName=self.brgParam1.get()
            tempSpeed=self.brgParam2.get()
            speed=1
            if is_number(tempSpeed)==False:
                speed=1
            else:
                speed=int(tempSpeed)/60.0
            if bearingName!='':
                for row in self.myTable.get_children():
                    self.myTable.delete(row)
                with self.con:
                    cur=self.con.cursor()
                    cur.execute(f"SELECT DISTINCT Brg_Number, BPFO, BPFI, BSF, FTF FROM Brg_Freqs WHERE Brg_Number = ? ORDER BY BPFO, BPFI, BSF, FTF", (bearingName,))
                    load_data = cur.fetchall()
                    load_data_arr = [i for i in load_data]
                    for i in range(len(load_data_arr)):
                        self.myTable.insert(parent='',index='end',iid=i,text='', values=(str(load_data_arr[i][0]),str(load_data_arr[i][1]*speed)[:6], \
                                            str(load_data_arr[i][2]*speed)[:6], str(load_data_arr[i][3]*speed)[:6], str(load_data_arr[i][4]*speed)[:6]))
                    
        except:
            pass
class SideButtonFrame(Tk.Frame):
    def __init__(self, parent, history_config_struct, db_connect, infoLabel, canvas):
        super().__init__(parent, bd=1, bg='white', width=90, height=504)
        self.parent=parent
        self.link="https://dantri.com.vn/"
        self.history_config_struct=history_config_struct
        self.infoLabel=infoLabel
        self.canvas=canvas
        self.con=db_connect
        self.style = ttk.Style()
        self.style.configure('custom.TLabel', font=('Chakra Petch', 13), bg='white')
        self.style.configure('red.TLabel', font=('Chakra Petch', 13), bg='white', foreground='#C40069')
        self.style.configure('custom.TLabelframe', font=('Chakra Petch', 15), bg='white', borderwidth=0)
        self.style.configure('custom.Accent.TButton', font=('Chakra Petch', 10), justify=Tk.CENTER)
        self.style.configure('small.Accent.TButton', font=('Chakra Petch', 9), justify=Tk.CENTER)
        imageAddress = ImageAdrr()
        self.ZoomCanvas = Tk.Canvas()
        self.functionCanvas = Tk.Canvas()
        self.zoomPhoto = imageAddress.zoomPhoto
        self.savePhoto = imageAddress.savePhoto
        self.zoomIn = imageAddress.zoomIn
        self.zoomOut = imageAddress.zoomOut
        self.panLeft = imageAddress.panLeft
        self.panRight = imageAddress.panRight
        self.function1 = imageAddress.fuction1
        self.creat_button()

    def creat_button(self):
        saveBt = ttk.Button(self, style='custom.Accent.TButton', text=_("EXPORT\nREPORT"), command=self.report)
        saveBt.place(x=0, y=425, width=88, height=75)

        exportDataBt = ttk.Button(self, style='custom.Accent.TButton', text=_("EXPORT\n.CSV"), command=self.export_from_db)
        exportDataBt.place(x=0, y=348, width=88, height=75)

        freqZoomBt = ttk.Button(self, style='custom.Accent.TButton', text=_("ZOOM"),
                                image=self.zoomPhoto,
                                compound=Tk.TOP,
                                command=lambda: self.on_zoom_button_clicked(827, 117))
        freqZoomBt.place(x=0, y=271, width=88, height=75)
        freqZoomBt.image = self.zoomPhoto

        self.freqCursorLeftBt = ttk.Button(self, style='custom.Accent.TButton', text=_("CURSOR\nLEFT"),
                                command=lambda:self.Tracking(False))
        self.freqCursorLeftBt.place(x=0, y=194, width=88, height=75)

        self.freqCursorRightBt = ttk.Button(self, style='custom.Accent.TButton', text=_("CURSOR\nRIGHT"),
                                command=lambda:self.Tracking(True))
        self.freqCursorRightBt.place(x=0, y=117, width=88, height=75)

        # self.freqGridtBt = ttk.Button(self, style='custom.Accent.TButton', text="RESERVE", state="disable")
        # self.freqGridtBt.place(x=0, y=117, width=88, height=75)

        self.freqFunctionBt = ttk.Button(self, style='custom.Accent.TButton', text=_("PLOT\nNONE"),
                                command=lambda:self.on_plot_button_clicked(827, 40))
        self.freqFunctionBt.place(x=0, y=40, width=88, height=75)

    def on_zoom_button_clicked(self, x_pos, y_pos):
        global blink, blink1
        if blink1 == 1:
            self.functionCanvas.destroy()
            blink1=0
        blink = not blink
        if blink == 1:
            self.creat_zoom_button_canvas(self.parent.historyPlotFrame, self.canvas, x_pos,
                                                y_pos)
        elif blink == 0:
            self.ZoomCanvas.destroy()
       
    def creat_zoom_button_canvas(self, widget, draw_canvas, x_pos, y_pos):

        self.zoomstyle=ttk.Style()
        self.zoomstyle.configure('zoom.Accent.TButton', font=('Chakra Petch', 9), justify=Tk.CENTER)
        self.ZoomCanvas = Tk.Canvas(widget, width=90, height=385, bg='white')
        self.ZoomCanvas.place(x=x_pos, y=y_pos)

        button1 = ttk.Button(self.ZoomCanvas, text=_("ZOOM IN"), style='zoom.Accent.TButton', image=self.zoomIn,
                             compound=Tk.TOP,
                             command=lambda: self.view_change(draw_canvas, _type='IN'))
        button1.place(x=0, y=0, width=88, height=75)
        button1.image = self.zoomIn

        button2 = ttk.Button(self.ZoomCanvas, text=_("ZOOM OUT"), style='zoom.Accent.TButton', image=self.zoomOut,
                             compound=Tk.TOP,
                             command=lambda: self.view_change(draw_canvas, _type='OUT'))
        button2.place(x=0, y=77, width=88, height=75)
        button2.image = self.zoomOut
        button3 = ttk.Button(self.ZoomCanvas, text=_("PAN LEFT"), style='zoom.Accent.TButton', image=self.panLeft,
                             compound=Tk.TOP,
                             command=lambda: self.view_change(draw_canvas, _type='LEFT'))
        button3.place(x=0, y=154, width=88, height=75)
        button3.image = self.panLeft

        button4 = ttk.Button(self.ZoomCanvas, text=_("PAN RIGHT"), style='zoom.Accent.TButton', image=self.panRight,
                             compound=Tk.TOP,
                             command=lambda: self.view_change(draw_canvas, _type='RIGHT'))
        button4.place(x=0, y=231, width=88, height=75)
        button4.image = self.panRight

        button5 = ttk.Button(self.ZoomCanvas, text=_("RESET"), style='zoom.Accent.TButton',
                            command=lambda: self.view_change(draw_canvas, _type='RESET'))
        button5.place(x=0, y=308, width=88, height=75)

    def view_change(self, canvas, _type):
        axes_arr = canvas.figure.get_axes()
        try:
            for ax in axes_arr:
                if _type == "RIGHT":
                    [xleft, xright] = ax.get_xlim()
                    xright -= 0.05*(xright-xleft)
                    xleft -= 0.05*(xright-xleft)
                    ax.set_xlim(xleft, xright)

                elif _type == "LEFT":
                    [xleft, xright] = ax.get_xlim()
                    xright += 0.05*(xright-xleft)
                    xleft += 0.05*(xright-xleft)
                    if xright >= 0:
                        ax.set_xlim(xleft, xright)

                elif _type == "IN":
                    [xleft, xright] = ax.get_xlim()
                    xright -= 0.05*(xright-xleft)
                    if xright < xleft*1.1:
                        xright = xleft*1.1
                    ax.set_xlim(xleft, xright)

                elif _type == "OUT":
                    [xleft, xright] = ax.get_xlim()
                    xright += 0.05*(xright-xleft)
                    ax.set_xlim(xleft, xright)
                elif _type == "RESET":
                    [xleft, xright] = [0, self.history_config_struct["ViewLimit"]]
                    ax.set_xlim(xleft, xright)
            canvas.draw()
        except:
            pass
    def on_plot_button_clicked(self, x_pos, y_pos):
        global blink, blink1
        if blink == 1:
            self.ZoomCanvas.destroy()
            blink=0
        blink1 = not blink1
        if blink1 == 1:
            self.creat_plot_button_canvas(self.parent.historyPlotFrame, self.canvas, x_pos,
                                                y_pos)
        elif blink1 == 0:
            self.functionCanvas.destroy()
       
    def creat_plot_button_canvas(self, widget, draw_canvas, x_pos, y_pos):

        self.plotStyle=ttk.Style()
        self.plotStyle.configure('plot.Accent.TButton', font=('Chakra Petch', 9), justify=Tk.CENTER)
        self.functionCanvas = Tk.Canvas(widget, width=90, height=462, bg='white')
        self.functionCanvas.place(x=x_pos, y=y_pos)

        waveformBt = ttk.Button(self.functionCanvas, text=_("WAVEFORM"), style='plot.Accent.TButton',
                             command=lambda: self.history_plot("WAVEFORM"))
        waveformBt.place(x=0, y=0, width=88, height=75)

        frequencyBt = ttk.Button(self.functionCanvas, text=_("FREQUENCY"), style='plot.Accent.TButton',
                             command=lambda: self.history_plot('FREQUENCY'))
        frequencyBt.place(x=0, y=77, width=88, height=75)

        velocityBt = ttk.Button(self.functionCanvas, text=_("VELOCITY\nFREQUENCY"), style='plot.Accent.TButton',
                             command=lambda: self.history_plot('VEL FREQUENCY'))
        velocityBt.place(x=0, y=154, width=88, height=75)

        envelopBt = ttk.Button(self.functionCanvas, text=_("ENVELOPED"), style='plot.Accent.TButton',
                             command=lambda: self.history_plot('ENVELOPED'))
        envelopBt.place(x=0, y=231, width=88, height=75)

        waterFallBt = ttk.Button(self.functionCanvas, text=_("WATERFALL"), style='plot.Accent.TButton',
                                command=lambda: self.history_plot('WATERFALL'))
        waterFallBt.place(x=0, y=308, width=88, height=75)

        trendBt = ttk.Button(self.functionCanvas, text=_("TREND"), style='plot.Accent.TButton',
                                command=lambda: self.history_plot('TREND'))
        trendBt.place(x=0, y=385, width=88, height=75)

    def history_plot(self, plot_type:str, export_png=False):
        global track_flag
        track_flag=0
        self.freqCursorLeftBt.configure(state='normal')
        self.freqCursorRightBt.configure(state='normal')
        self.freqCursorLeftBt.update_idletasks()
        self.freqCursorRightBt.update_idletasks()
        try:
            prjCode = self.history_config_struct["ProjectID"]
            sensorPosition = self.history_config_struct["SensorPosition"]
            codeTuple=(prjCode,)
            dataTuple=(prjCode, sensorPosition,)
            viewRange=[0, self.history_config_struct["ViewLimit"]]
            tsaUse = self.history_config_struct["TSA"]
            tsaBin = self.history_config_struct["TsaBin"]
            result_arr1=[]
            date_arr=[]
            sample_rate_arr=[]
            with self.con:
                cur=self.con.cursor()
                cur.execute(f"SELECT DATA, DATE, Sample_rate FROM DATA WHERE CODE = ? AND POS=? ORDER BY DATE DESC;", dataTuple)
                load_data = cur.fetchall()
                cur.execute(f"SELECT POWER, RPM, DRIVEN, BEARINGBORE, GEARTOOTH, FOUNDATION, NOTE FROM Project_ID WHERE CODE = ?", codeTuple)
                machineParam = cur.fetchall()
            load_data_arr = [i for i in load_data] #mang cac chuoi data, date [data,date,sample_rate; data, date, sample_rate]
            machine_param_arr = [i for i in machineParam]
            if len(load_data_arr)!=0:
                for ari in load_data_arr:
                    result_arr1.append(file_operation.extract_str(ari[0]))
                    date_arr.append(ari[1])
                    sample_rate_arr.append(ari[2])
                if len(result_arr1)>6:
                    result_arr2=result_arr1[0:6]
                else:
                    result_arr2=result_arr1
                for ari in machine_param_arr:
                    power=ari[0]
                    rpm=ari[1]
                    driven=ari[2]
                    bearing_bore=ari[3]
                    gear_tooth=ari[4]
                    foundation=ari[5]
                    note=ari[6]
                self.history_config_struct["dataSample"]=result_arr2
                self.history_config_struct["sampleRate"]=sample_rate_arr
                if plot_type=="WAVEFORM":
                    self.freqFunctionBt.configure(text=_("PLOT\nWAVEFORM"), style="small.Accent.TButton")
                    Pd.PLT.plot_all_history(self.canvas, result_arr2, date_arr, sample_rate_arr, viewRange, 1, tsaUse, tsaBin, export_png=export_png)
                elif plot_type=="FREQUENCY":
                    self.freqFunctionBt.configure(text=_("PLOT\nFREQUENCY\nSPECTRUM"), style="small.Accent.TButton")
                    Pd.PLT.plot_all_history(self.canvas, result_arr2, date_arr, sample_rate_arr, viewRange, 0, tsaUse, tsaBin, export_png=export_png)
                elif plot_type=="VEL FREQUENCY":
                    if sensorPosition[-1]=='A':
                        self.freqFunctionBt.configure(text=_("PLOT\nVELOCITY\nSPECTRUM"), style="small.Accent.TButton")
                        # Pd.PLT.plot_velocity_spectral_kiem_tra_tich_phan(self.canvas, result_arr2, date_arr, sample_rate_arr, viewRange, export_png=export_png)
                        Pd.PLT.plot_history_velocity_spectral(self.canvas, result_arr2, date_arr, sample_rate_arr, viewRange, tsaUse, tsaBin, export_png=export_png)

                elif plot_type=="WATERFALL":
                    self.freqFunctionBt.configure(text=_("PLOT\nWATERFALL"), style="small.Accent.TButton")
                    Pd.PLT.plot_waterfall(self.canvas, result_arr1, date_arr, sample_rate_arr, viewRange, export_png=export_png)

                elif plot_type=="TREND":
                    self.freqFunctionBt.configure(text=_("PLOT\nTREND"), style="small.Accent.TButton")
                    self.freqCursorLeftBt.configure(state='disable')
                    self.freqCursorRightBt.configure(state='disable')
                    self.freqCursorLeftBt.update_idletasks()
                    self.freqCursorRightBt.update_idletasks()
                    view_arr=[]
                    view_arr.append(self.history_config_struct["Rms"])
                    view_arr.append(self.history_config_struct["A_Pk"])
                    view_arr.append(self.history_config_struct["gE"])
                    view_arr.append(self.history_config_struct["HFCF"])
                    isoStandard = iso10816_judge(driven, rpm, power, foundation)
                    Pd.PLT.plot_trend(self.canvas, result_arr1, date_arr, sample_rate_arr, sensorPosition,\
                    isoStandard, rpm, bearing_bore, view_arr, export_png=export_png)
                elif plot_type=="ENVELOPED":
                    self.freqFunctionBt.configure(text=_("PLOT\nENVELOPED"), style="small.Accent.TButton")
                    filter_type = "BANDPASS"
                    filter_from = self.history_config_struct["FilterFrom"]# high pass cutoff freq
                    filter_to = self.history_config_struct["FilterTo"]# low pass cutoff freq
                    
                    if filter_from >= filter_to:
                        filter_from= dfc._ENV_BANDPASS_FROM
                        filter_to= dfc._ENV_BANDPASS_TO
                        self.infoLabel.config(text=_("Fcut is wrong. Use the default value."))
                    else:
                        pass
                    temp_result_arr=[]
                    if filter_type:
                        for i in range(len(result_arr2)):
                            _samples_1=filter_data(
                                result_arr2[i],
                                filter_type,
                                filter_from,
                                filter_to,
                                sample_rate_arr[i],
                                window="Hanning"
                                )
                            analytical_signal1 = hilbert(_samples_1)
                            # enveloped_signal1=_samples_1 + 1j*analytical_signal1
                            amplitude_envelope1 = np.abs(analytical_signal1)
                            amplitude_envelope1 = filter_data(
                                amplitude_envelope1,
                                "LOWPASS",
                                filter_from,
                                filter_to/2,
                                sample_rate_arr[i],
                                window="Hanning"
                                )
                            temp_result_arr.append(amplitude_envelope1)
                    Pd.PLT.plot_all_history(self.canvas, temp_result_arr, date_arr, sample_rate_arr, viewRange, 2, tsaUse, tsaBin, export_png=export_png)
                self.infoLabel.config(text=_("Project: ")+ str(prjCode) + '. ')
            else:
                self.infoLabel.config(text=_("There is no data, please check SETTING !"))
        except Exception as ex:
            print(ex)
            self.infoLabel.config(text=_("Data errors."))

    def Tracking(self, dir:bool):
        global track_flag
        tracking_freq=self.history_config_struct["TrackingResolution"]
        axes_arr = self.canvas.figure.get_axes()     
                       
        if len(axes_arr)>0:
            if axes_arr[0].name=='3d':
                return
            else:
                x_data=axes_arr[0].get_lines()[-1].get_xdata()
                y_data=axes_arr[0].get_lines()[-1].get_ydata()

            [xleft, xright] = axes_arr[0].get_xlim()
            
            if dir==True:
                if xleft>=track_flag:
                    track_flag=xleft
                if xright<track_flag:
                    return
                track_flag+=tracking_freq
                start_freq=track_flag-tracking_freq
                stop_freq=track_flag
            else:
                if xright<=track_flag:
                    track_flag = xright

                track_flag-=tracking_freq

                if xleft>=track_flag:
                    track_flag = xleft
                if track_flag<=0:
                    track_flag=0
                start_freq=track_flag
                stop_freq=track_flag+tracking_freq
            try:     
                _sample_rate=self.history_config_struct["sampleRate"][0]
                [max1, freq]=tab4_tracking_signal(y_data, x_data, [start_freq, stop_freq])
                Pd.PLT.plot_grid_only(self.canvas, freq)
                title= _('Frequency:')+ f' {str(freq)[:4]}'+' hz'
                self.infoLabel.config(text=title)
            except:
                self.infoLabel.config(text=_("Data errors."))


    def export_from_db(self):
        try:
            if pms.general_warning(_("Do you want to export data to .CSV file")):
                result_arr2=[]
                pos_arr=[]
                date_arr=[]
                
                file_name = ''
                save_path=f"{parent_directory}/storage/"
                ProjectCode=self.history_config_struct["ProjectID"]
                
                with self.con:
                    cur = self.con.cursor()
                    cur.execute(f"SELECT DATA, POS, DATE FROM DATA WHERE CODE = ? ORDER BY DATE ASC", (ProjectCode,))
                    load_data = cur.fetchall()
                    cur.execute(f"SELECT COM_ID FROM Project_ID WHERE CODE = ?", (ProjectCode,))
                    machineParam = cur.fetchall()
                load_data_arr = [i for i in load_data]
                machine_param_arr = [i for i in machineParam]
                comId=machine_param_arr[0][0]
                
                with self.con:
                    cur = self.con.cursor()
                    cur.execute(f"SELECT NAME FROM Company_ID WHERE COM_ID = ?", (str(comId),))
                    companyTuple=cur.fetchall()
                companyName= companyTuple[0][0]
                name=companyName + '_' + ProjectCode+'__'
                if name =='':
                    file_name +='NONAME_DATA__-__-__'
                else:
                    file_name += name

                if len(load_data_arr)!=0:
                    for ari in load_data_arr:
                        result_arr2.append(file_operation.extract_str(ari[0]))
                        pos_arr.append(ari[1])
                        date_arr.append(ari[2])
                    

                    for i in range(len(pos_arr)):
                        new_file_name= f'_{str(i+1)}_' + file_name + date_arr[i] + pos_arr[i] + '.csv'
                        completeName = os.path.join(save_path, new_file_name)
                        file_operation.store_data(result_arr2[i], completeName)
                        
                    
                    try:
                        os.system("sudo umount /media/pi/usb")
                    except:
                        pass
                    try:
                        os.system("sudo rmdir /media/pi/usb")
                    except:
                        pass
                    try:
                        os.system("sudo mkdir /media/pi/usb")
                    except:
                        pass
                    try:
                        os.system("sudo mount /dev/sda1 /media/pi/usb")
                    except:
                        pass
                    try:
                        os.system(f"sudo cp -ru -f {save_path}/* /media/pi/usb" )
                        os.system("sudo umount /media/pi/usb")
                        os.system("sudo rmdir /media/pi/usb")
                        self.infoLabel.configure(text=_("Export CSV is completed"))
                    except Exception as ex:
                        print(ex)
                    try:
                        os.system(f"sudo rm -r {save_path}/*")
                    except Exception as ex:
                        print(ex)
                else:
                    self.infoLabel.configure(text=_("There is no data, please check input value !"))
            else:
                pass
        except Exception as ex:
            print("Export from db error", ex)


    def report(self):
        from docx.shared import Inches
        save_path=f"{parent_directory}/storage/"
        document=rp.Report(save_path)
        try:
            if pms.general_warning(_("Do you want to export the report")):
                prjCode = self.history_config_struct["ProjectID"]
                sensorPosition = self.history_config_struct["SensorPosition"]
                viewRange=[0, self.history_config_struct["ViewLimit"]]
                tsaUse = self.history_config_struct["TSA"]
                tsaBin = self.history_config_struct["TsaBin"]
                result_arr1=[]
                date_arr=[]
                sample_rate_arr=[]
                with self.con:
                    cur=self.con.cursor()
                    cur.execute(f"SELECT DATA, DATE, Sample_rate FROM DATA WHERE CODE = ? AND POS= ? ORDER BY DATE DESC", (prjCode, sensorPosition,))
                    load_data = cur.fetchall()
                    cur.execute(f"SELECT POWER, RPM, DRIVEN, BEARINGBORE, GEARTOOTH, FOUNDATION, NOTE FROM Project_ID WHERE CODE = ?", (prjCode,))
                    machineParam = cur.fetchall()
                load_data_arr = [i for i in load_data] #mang cac chuoi data, date [data,date,sample_rate; data, date, sample_rate]
                machine_param_arr = [i for i in machineParam]
                if len(load_data_arr)!=0:
                    for ari in load_data_arr:
                        result_arr1.append(file_operation.extract_str(ari[0]))
                        date_arr.append(ari[1])
                        sample_rate_arr.append(ari[2])
                    if len(result_arr1)>6:
                        result_arr2=result_arr1[0:6]
                    else:
                        result_arr2=result_arr1
                    for ari in machine_param_arr:
                        power=ari[0]
                        rpm=ari[1]
                        driven=ari[2]
                        bearing_bore=ari[3]
                        gear_tooth=ari[4]
                        foundation=ari[5]
                        note=ari[6]
                    # Export waveform image
                    Pd.PLT.plot_all_history(self.canvas, result_arr2, date_arr, sample_rate_arr, viewRange, 1, tsaUse, tsaBin, export_png=True)
                    
                    # Export acc frequency image
                    Pd.PLT.plot_all_history(self.canvas, result_arr2, date_arr, sample_rate_arr, viewRange, 0, tsaUse, tsaBin, export_png=True)
                    
                    # Export velocity frequency image
                    if sensorPosition[-1]=='A':
                        Pd.PLT.plot_history_velocity_spectral(self.canvas, result_arr2, date_arr, sample_rate_arr, [0, 1000], tsaUse, tsaBin, export_png=True)
                    
                    # Export waterfall image  
                    Pd.PLT.plot_waterfall(self.canvas, result_arr1, date_arr, sample_rate_arr, viewRange, export_png=True)

                    # Export Iso image
                    view_arr=[]
                    view_arr.append(self.history_config_struct["Rms"])
                    view_arr.append(self.history_config_struct["A_Pk"])
                    view_arr.append(self.history_config_struct["gE"])
                    view_arr.append(self.history_config_struct["HFCF"])
                    isoStandard = iso10816_judge(driven, rpm, power, foundation)
                    Pd.PLT.plot_trend(self.canvas, result_arr1, date_arr, sample_rate_arr, sensorPosition,\
                    isoStandard, rpm, bearing_bore, view_arr, export_png=True)

                    # Export envelop image
                    filter_type = "BANDPASS"
                    filter_from = self.history_config_struct["FilterFrom"]# high pass cutoff freq
                    filter_to = self.history_config_struct["FilterTo"]# low pass cutoff freq
                    
                    if filter_from >= filter_to:
                        filter_from= dfc._ENV_BANDPASS_FROM
                        filter_to= dfc._ENV_BANDPASS_TO
                        self.infoLabel.configure(text=_("Filter bandpass frequency is wrong. Use the default value."))
                    else:
                        pass
                    temp_result_arr=[]
                    if filter_type:
                        for i in range(len(result_arr2)):
                            _samples_1=filter_data(
                                result_arr2[i],
                                filter_type,
                                filter_from,
                                filter_to,
                                sample_rate_arr[i],
                                window="Hanning"
                                )
                            analytical_signal1 = hilbert(_samples_1)
                            # enveloped_signal1=_samples_1 + 1j*analytical_signal1
                            amplitude_envelope1 = np.abs(analytical_signal1)
                            amplitude_envelope1 = filter_data(
                                amplitude_envelope1,
                                "LOWPASS",
                                filter_from,
                                filter_to/2,
                                sample_rate_arr[i],
                                window="Hanning"
                                )
                            temp_result_arr.append(amplitude_envelope1)
                    Pd.PLT.plot_all_history(self.canvas, temp_result_arr, date_arr, sample_rate_arr, [0, 2000], 2, tsaUse, tsaBin, export_png=True)
                    document.first_page()
                    document._add_picture(save_path+'trend.png', width=Inches(6.0))
                    document._add_run(_('Figure1: Vibration trend'), style='italic')
                    document.add_blank_comment()

                    document._add_picture(save_path+'waveform.png', width=Inches(6.0))
                    document._add_run(_('Figure2: Waveform'), style='italic')
                    document.add_blank_comment()

                    document._add_picture(save_path+'frequency.png', width=Inches(6.0))
                    document._add_run(_('Figure3: Frequency spectrum'), style='italic')
                    document.add_blank_comment()

                    document._add_picture(save_path+'velocity_frequency.png', width=Inches(6.0))
                    document._add_run(_('Figure4: Velocity spectrum'), style='italic')
                    document.add_blank_comment()

                    document._add_picture(save_path+'envelop.png', width=Inches(6.0))
                    document._add_run(_('Figure5: Envelope spectrum'), style='italic')
                    document.add_blank_comment()
                    document._add_page_break()
                    new_file_name=str(prjCode)+'_report.docx'
                    document._save(new_file_name)

                    try:
                        os.system("sudo umount /media/pi/usb")
                    except:
                        pass
                    try:
                        os.system("sudo rmdir /media/pi/usb")
                    except:
                        pass
                    try:
                        os.system("sudo mkdir /media/pi/usb")
                    except:
                        pass
                    try:
                        os.system("sudo mount /dev/sda1 /media/pi/usb")
                    except:
                        pass
                    try:
                        os.system(f"sudo cp -ru -f {save_path}/{new_file_name} /media/pi/usb")
                        os.system("sudo umount /media/pi/usb")
                        os.system("sudo rmdir /media/pi/usb")
                        self.infoLabel.configure(text=_("Export CSV is completed"))
                    except Exception as ex:
                        print(ex)
                    try:
                        os.system(f"sudo rm -r {save_path}/*")
                    except Exception as ex:
                        print(ex)
                    # qr = qrcode.QRCode(version = 1, box_size = 10, border = 5)
                    # qr.add_data(self.link)
                    # qr.make(fit = True)
                    # img = qr.make_image(fill_color = 'black', back_color = 'white')
                    # img.save(save_path + 'MyQRCode2.png')
                    # Pd.PLT.plot_image(self.canvas, 'MyQRCode2.png', self.link)

                    self.infoLabel.configure(text=_("Report is exported. Use QR code scanner to see the report."))
                else:
                    self.infoLabel.configure(text=_("There is no data, please check SETTING !"))
        except:
            self.infoLabel.configure(text=_("Data errors."))

class HistoryPlotCanvas(Tk.Frame):
    def __init__(self, parent):
        super().__init__(parent, bd=1, bg='white', width=918, height=504)
        fig4 = plt.figure(figsize=(8.2, 6.1))
        fig4.set_visible(False)
        self.canvas4 = FigureCanvasTkAgg(fig4, master=self)
        self.canvas4.get_tk_widget().pack(side=Tk.TOP, fill=Tk.BOTH, expand=1)